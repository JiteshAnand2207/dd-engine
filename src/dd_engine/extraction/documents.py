"""Deterministic DOCX, CSV and standalone-image extractors."""

from __future__ import annotations

import csv
import hashlib
import io
from pathlib import Path

from docx import Document
from PIL import ExifTags, Image

from dd_engine.artifacts import atomic_write_bytes
from dd_engine.extraction.models import JsonObject, SourceExtraction, make_unit


def _normalize_text(value: str) -> str:
    return value.replace("\r\n", "\n").replace("\r", "\n").strip()


def _safe_suffix(content_type: str) -> str:
    suffixes = {
        "image/bmp": ".bmp",
        "image/gif": ".gif",
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/tiff": ".tiff",
        "image/webp": ".webp",
    }
    return suffixes.get(content_type.casefold(), ".bin")


def extract_docx(
    *,
    payload: bytes,
    source: JsonObject,
    run_id: str,
    run_path: Path,
    config_namespace: str,
) -> SourceExtraction:
    """Extract body paragraphs, tables and embedded-image metadata from DOCX."""

    document = Document(io.BytesIO(payload))
    units: list[JsonObject] = []
    warnings: list[str] = []
    current_heading: JsonObject | None = None
    heading_count = 0

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = _normalize_text(paragraph.text)
        style_name = paragraph.style.name if paragraph.style is not None else None
        is_heading = bool(style_name and style_name.casefold().startswith("heading"))
        heading_level: int | None = None
        if is_heading:
            heading_count += 1
            tail = style_name.rsplit(" ", 1)[-1] if style_name else ""
            heading_level = int(tail) if tail.isdigit() else None
            current_heading = {
                "heading_index": heading_count,
                "heading_level": heading_level,
                "paragraph_index": paragraph_index,
                "text": text,
            }
        if not text:
            continue
        locator: JsonObject = {
            "heading": current_heading,
            "heading_level": heading_level if is_heading else None,
            "paragraph_index": paragraph_index,
            "type": "docx_paragraph",
        }
        units.append(
            make_unit(
                run_id=run_id,
                source=source,
                ordinal=len(units) + 1,
                unit_type="docx_heading" if is_heading else "docx_paragraph",
                locator=locator,
                extraction_method="python_docx_paragraph",
                confidence=0.99,
                content={"style": style_name, "text": text},
            )
        )

    table_cell_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = _normalize_text("\n".join(item.text for item in cell.paragraphs))
                table_cell_count += 1
                units.append(
                    make_unit(
                        run_id=run_id,
                        source=source,
                        ordinal=len(units) + 1,
                        unit_type="docx_table_cell",
                        locator={
                            "cell_index": cell_index,
                            "cell_reference": f"R{row_index}C{cell_index}",
                            "row_index": row_index,
                            "table_index": table_index,
                            "type": "docx_table_cell",
                        },
                        extraction_method="python_docx_table",
                        confidence=0.99,
                        content={"paragraph_count": len(cell.paragraphs), "text": text},
                    )
                )

    embedded_images = 0
    seen_parts: set[str] = set()
    related_parts = getattr(document.part, "related_parts", {})
    for relationship_id, part in sorted(related_parts.items()):
        content_type = str(getattr(part, "content_type", ""))
        if not content_type.casefold().startswith("image/"):
            continue
        part_name = str(getattr(part, "partname", relationship_id))
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        blob = bytes(getattr(part, "blob", b""))
        embedded_images += 1
        suffix = _safe_suffix(content_type)
        relative_asset = (
            Path("extracts")
            / "cache"
            / "assets"
            / config_namespace
            / str(source["source_id"])
            / f"docx-image-{embedded_images:04d}{suffix}"
        )
        asset_path = run_path / relative_asset
        atomic_write_bytes(asset_path, blob)
        checksum = hashlib.sha256(blob).hexdigest()
        image_metadata: JsonObject = {
            "asset_checksum": checksum,
            "asset_path": relative_asset.as_posix(),
            "content_type": content_type,
            "size_bytes": len(blob),
        }
        try:
            with Image.open(io.BytesIO(blob)) as image:
                image_metadata.update(
                    {"height_pixels": image.height, "width_pixels": image.width}
                )
        except Exception as exc:
            image_metadata["metadata_warning"] = f"embedded image metadata unreadable: {exc}"
            warnings.append(str(image_metadata["metadata_warning"]))
        units.append(
            make_unit(
                run_id=run_id,
                source=source,
                ordinal=len(units) + 1,
                unit_type="docx_embedded_image",
                locator={
                    "image_number": embedded_images,
                    "relationship_id": relationship_id,
                    "type": "docx_image",
                },
                extraction_method="python_docx_embedded_image",
                confidence=0.95,
                content=image_metadata,
                warnings=(
                    [str(image_metadata["metadata_warning"])]
                    if "metadata_warning" in image_metadata
                    else []
                ),
            )
        )

    return SourceExtraction(
        status="successfully_extracted",
        primary_method="python_docx",
        units=units,
        warnings=list(dict.fromkeys(warnings)),
        metrics={
            "docx_embedded_images": embedded_images,
            "docx_headings": heading_count,
            "docx_paragraph_units": sum(
                unit["unit_type"] in {"docx_heading", "docx_paragraph"} for unit in units
            ),
            "docx_table_cells": table_cell_count,
            "docx_tables": len(document.tables),
        },
    )


def _column_letter(number: int) -> str:
    result = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        result = chr(65 + remainder) + result
    return result


def extract_csv(*, payload: bytes, source: JsonObject, run_id: str) -> SourceExtraction:
    """Extract actual CSV bytes regardless of the filename extension."""

    try:
        text = payload.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError(f"CSV is not valid UTF-8/UTF-8-SIG: {exc}") from exc
    sample = text[:65536]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text, newline=""), dialect))
    units: list[JsonObject] = []
    headers = rows[0] if rows else []
    for row_number, row in enumerate(rows, start=1):
        for column_number, value in enumerate(row, start=1):
            column_letter = _column_letter(column_number)
            units.append(
                make_unit(
                    run_id=run_id,
                    source=source,
                    ordinal=len(units) + 1,
                    unit_type="csv_cell",
                    locator={
                        "cell": f"{column_letter}{row_number}",
                        "column_index": column_number,
                        "column_letter": column_letter,
                        "column_name": (
                            headers[column_number - 1]
                            if row_number > 1 and column_number <= len(headers)
                            else None
                        ),
                        "row_index": row_number,
                        "type": "csv_cell",
                    },
                    extraction_method="python_csv",
                    confidence=0.99,
                    content={"value": value, "value_type": "string"},
                )
            )
    return SourceExtraction(
        status="successfully_extracted",
        primary_method="python_csv",
        units=units,
        metrics={
            "csv_cells": len(units),
            "csv_columns_max": max((len(row) for row in rows), default=0),
            "csv_delimiter": dialect.delimiter,
            "csv_rows": len(rows),
        },
    )


def _json_safe_exif(value: object) -> object:
    if value is None or isinstance(value, str | int | float | bool):
        return value
    if isinstance(value, bytes):
        return {"byte_count": len(value), "sha256": hashlib.sha256(value).hexdigest()}
    if isinstance(value, tuple):
        return [_json_safe_exif(item) for item in value]
    return str(value)


def extract_image(
    *,
    payload: bytes,
    source: JsonObject,
    run_id: str,
    run_path: Path,
    config_namespace: str,
) -> SourceExtraction:
    """Extract metadata and prepare one normalized local image for vision review."""

    with Image.open(io.BytesIO(payload)) as image:
        image.load()
        exif: JsonObject = {}
        try:
            for key, value in image.getexif().items():
                name = ExifTags.TAGS.get(key, str(key))
                exif[str(name)] = _json_safe_exif(value)
        except Exception:
            exif = {}
        metadata: JsonObject = {
            "dpi": list(image.info.get("dpi", ())) or None,
            "exif": exif,
            "format": image.format,
            "frame_count": int(getattr(image, "n_frames", 1)),
            "height_pixels": image.height,
            "mode": image.mode,
            "width_pixels": image.width,
        }
        normalized = image.convert("RGB") if image.mode not in {"RGB", "RGBA"} else image.copy()
        buffer = io.BytesIO()
        normalized.save(buffer, format="PNG")
    rendered = buffer.getvalue()
    relative_asset = (
        Path("extracts")
        / "rendered_pages"
        / config_namespace
        / str(source["source_id"])
        / "image-0001.png"
    )
    atomic_write_bytes(run_path / relative_asset, rendered)
    asset_checksum = hashlib.sha256(rendered).hexdigest()
    unit = make_unit(
        run_id=run_id,
        source=source,
        ordinal=1,
        unit_type="image_metadata",
        locator={
            "image_number": 1,
            "region": {
                "height": metadata["height_pixels"],
                "width": metadata["width_pixels"],
                "x": 0,
                "y": 0,
            },
            "type": "image",
        },
        extraction_method="pillow_image_metadata",
        confidence=0.99,
        content=metadata,
        limitation="visual content has not been interpreted",
    )
    task: JsonObject = {
        "asset": {
            "height_pixels": metadata["height_pixels"],
            "mime_type": "image/png",
            "path": relative_asset.as_posix(),
            "sha256": asset_checksum,
            "width_pixels": metadata["width_pixels"],
        },
        "limitation": "deterministic extraction captured metadata only",
        "locator": unit["locator"],
        "model_result": None,
        "relative_path": source["relative_path"],
        "requested_route": "codex_or_claude_vision_review",
        "run_id": run_id,
        "source_checksum": source["sha256"],
        "source_id": source["source_id"],
        "status": "pending",
        "task_id": f"VISION-{source['source_id']}-IMAGE-0001",
        "untrusted_source_data": True,
    }
    return SourceExtraction(
        status="queued_for_vision",
        primary_method="pillow_image_metadata",
        units=[unit],
        vision_tasks=[task],
        limitation="image meaning requires pending local vision review",
        metrics={"image_metadata_units": 1, "standalone_images_queued": 1},
    )
