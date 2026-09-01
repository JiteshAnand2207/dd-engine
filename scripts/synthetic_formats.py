"""Deterministic document writers for the fictional Phase 3 data room.

The helpers deliberately avoid Office, LibreOffice, browser rendering and model
APIs.  They create ordinary local artifacts and normalize archive timestamps so
that a fixed seed produces byte-identical output on the supported dependency
versions.
"""

from __future__ import annotations

import io
import math
import random
import zipfile
from collections.abc import Iterable, Sequence
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape, quoteattr

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

FIXED_DT = datetime(2026, 8, 31, 9, 0, 0, tzinfo=UTC)
ZIP_DT = (2026, 8, 31, 9, 0, 0)
PAGE_W, PAGE_H = A4
SYNTHETIC_LABEL = "FICTIONAL SYNTHETIC DATA - NOT A REAL COMPANY"


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    name = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
    try:
        return ImageFont.truetype(name, size=size)
    except OSError:
        return ImageFont.load_default()


def _fixed_zip_info(name: str) -> zipfile.ZipInfo:
    info = zipfile.ZipInfo(name, ZIP_DT)
    info.compress_type = zipfile.ZIP_DEFLATED
    info.create_system = 0
    info.external_attr = 0o600 << 16
    return info


def write_stable_zip(path: Path, members: dict[str, bytes]) -> None:
    """Write a deterministic ZIP with sorted, file-only members."""

    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w") as archive:
        for name in sorted(members):
            archive.writestr(_fixed_zip_info(name.replace("\\", "/")), members[name])


def normalize_ooxml(path: Path) -> None:
    """Normalize ZIP member ordering and timestamps for DOCX/XLSX determinism."""

    with zipfile.ZipFile(path) as source:
        members = {name: source.read(name) for name in source.namelist()}
    write_stable_zip(path, members)


def _pdf_footer(pdf: canvas.Canvas, page_number: int, accent: tuple[float, float, float]) -> None:
    pdf.setStrokeColorRGB(*accent)
    pdf.setLineWidth(0.5)
    pdf.line(42, 34, PAGE_W - 42, 34)
    pdf.setFillColorRGB(0.32, 0.34, 0.38)
    pdf.setFont("Helvetica", 7.5)
    pdf.drawString(42, 22, SYNTHETIC_LABEL)
    pdf.drawRightString(PAGE_W - 42, 22, f"Page {page_number}")


def _wrapped_lines(text: str, width_chars: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        if current and len(" ".join(current + [word])) > width_chars:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines or [""]


def write_text_pdf(
    path: Path,
    *,
    title: str,
    subtitle: str,
    sections: Sequence[tuple[str, Sequence[str]]],
    table: Sequence[Sequence[Any]] | None = None,
    template: str = "navy",
    metadata: Sequence[tuple[str, str]] = (),
) -> None:
    """Create a polished, searchable A4 PDF using one of three templates."""

    palettes = {
        "navy": ((0.06, 0.17, 0.29), (0.82, 0.67, 0.28), "Helvetica"),
        "green": ((0.08, 0.32, 0.25), (0.66, 0.82, 0.72), "Helvetica"),
        "slate": ((0.22, 0.25, 0.29), (0.71, 0.76, 0.82), "Times-Roman"),
    }
    accent, secondary, body_font = palettes[template]
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(path), pagesize=A4, invariant=1, pageCompression=1, pdfVersion=(1, 5))
    pdf.setTitle(title)
    pdf.setAuthor("dd-engine fictional fixture generator")
    page = 1

    def new_page(*, continuation: bool = False) -> float:
        nonlocal page
        if page > 1:
            pdf.showPage()
        pdf.setFillColorRGB(*accent)
        pdf.rect(0, PAGE_H - 88, PAGE_W, 88, stroke=0, fill=1)
        pdf.setFillColorRGB(1, 1, 1)
        pdf.setFont("Helvetica-Bold", 17 if continuation else 21)
        heading = f"{title} - continued" if continuation else title
        pdf.drawString(42, PAGE_H - 50, heading[:80])
        pdf.setFont("Helvetica", 9)
        pdf.drawString(42, PAGE_H - 69, subtitle[:105])
        _pdf_footer(pdf, page, accent)
        page += 1
        return PAGE_H - 112

    y = new_page()
    if metadata:
        pdf.setFillColorRGB(0.96, 0.97, 0.98)
        pdf.roundRect(42, y - 54, PAGE_W - 84, 48, 4, stroke=0, fill=1)
        pdf.setFillColorRGB(0.12, 0.14, 0.17)
        pdf.setFont("Helvetica", 8.5)
        x = 52
        for label, value in metadata:
            pdf.setFont("Helvetica-Bold", 8)
            pdf.drawString(x, y - 23, label.upper())
            pdf.setFont("Helvetica", 8.5)
            pdf.drawString(x, y - 38, value[:30])
            x += (PAGE_W - 104) / max(1, len(metadata))
        y -= 72

    for section_heading, paragraphs in sections:
        if y < 105:
            y = new_page(continuation=True)
        pdf.setFillColorRGB(*accent)
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(42, y, section_heading)
        y -= 18
        pdf.setFillColorRGB(0.10, 0.11, 0.13)
        pdf.setFont(body_font, 9.2)
        for paragraph in paragraphs:
            for line in _wrapped_lines(paragraph, 104):
                if y < 58:
                    y = new_page(continuation=True)
                    pdf.setFillColorRGB(0.10, 0.11, 0.13)
                    pdf.setFont(body_font, 9.2)
                pdf.drawString(48, y, line)
                y -= 12.5
            y -= 5
        y -= 7

    if table:
        column_count = max(len(row) for row in table)
        usable = PAGE_W - 84
        widths = [
            usable * (0.40 if index == 0 else 0.60 / max(1, column_count - 1))
            for index in range(column_count)
        ]
        for row_number, row in enumerate(table):
            line_count = max(len(_wrapped_lines(str(value), 25)) for value in row)
            height = max(21, 10 + line_count * 10)
            if y - height < 48:
                y = new_page(continuation=True)
            x = 42
            for index in range(column_count):
                value = str(row[index]) if index < len(row) else ""
                if row_number == 0:
                    pdf.setFillColorRGB(*accent)
                    pdf.rect(x, y - height, widths[index], height, stroke=0, fill=1)
                    pdf.setFillColorRGB(1, 1, 1)
                    pdf.setFont("Helvetica-Bold", 8)
                else:
                    fill = 0.97 if row_number % 2 else 0.93
                    pdf.setFillColorRGB(fill, fill, fill)
                    pdf.rect(x, y - height, widths[index], height, stroke=0, fill=1)
                    pdf.setFillColorRGB(0.10, 0.11, 0.13)
                    pdf.setFont("Helvetica", 8)
                text_y = y - 13
                for line in _wrapped_lines(value, max(10, int(widths[index] / 5.2))):
                    pdf.drawString(x + 5, text_y, line)
                    text_y -= 9
                pdf.setStrokeColorRGB(*secondary)
                pdf.rect(x, y - height, widths[index], height, stroke=1, fill=0)
                x += widths[index]
            y -= height
    pdf.save()


def _draw_scan_page(
    *, title: str, lines: Sequence[str], seed: int, stamp: str | None = None
) -> Image.Image:
    rng = random.Random(seed)
    width, height = 1240, 1754
    image = Image.new("RGB", (width, height), (244, 242, 234))
    draw = ImageDraw.Draw(image)
    for _ in range(2500):
        shade = rng.randint(224, 250)
        x = rng.randrange(width)
        y = rng.randrange(height)
        draw.point((x, y), fill=(shade, shade, max(215, shade - 3)))
    draw.rectangle((82, 70, width - 82, height - 72), outline=(110, 108, 102), width=3)
    draw.text((110, 105), title, font=_font(36, bold=True), fill=(36, 36, 34))
    draw.text((110, 160), SYNTHETIC_LABEL, font=_font(16, bold=True), fill=(145, 35, 35))
    y = 225
    for paragraph in lines:
        for line in _wrapped_lines(paragraph, 76):
            draw.text((118, y), line, font=_font(22), fill=(55, 54, 50))
            y += 33
        y += 22
    if stamp:
        draw.rounded_rectangle((760, 1390, 1110, 1510), radius=14, outline=(150, 40, 40), width=6)
        draw.text((790, 1430), stamp, font=_font(28, bold=True), fill=(150, 40, 40))
    return image.filter(ImageFilter.GaussianBlur(radius=0.25))


def write_image_only_pdf(
    path: Path, *, title: str, page_lines: Sequence[Sequence[str]], seed: int
) -> None:
    """Create a PDF whose pages contain raster images and no text operators."""

    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(path), pagesize=A4, invariant=1, pageCompression=1)
    pdf.setTitle(title)
    for index, lines in enumerate(page_lines):
        page_image = _draw_scan_page(
            title=f"{title} - page {index + 1}",
            lines=lines,
            seed=seed + index,
            stamp="SCANNED COPY" if index == 0 else None,
        )
        payload = io.BytesIO()
        page_image.save(payload, format="JPEG", quality=72, optimize=False, progressive=False)
        payload.seek(0)
        pdf.drawImage(ImageReader(payload), 0, 0, PAGE_W, PAGE_H)
        pdf.showPage()
    pdf.save()


def write_corrupt_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"%PDF-1.7\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n")


def _cell_width(cell: Any, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def write_docx(
    path: Path,
    *,
    title: str,
    subtitle: str,
    metadata: Sequence[tuple[str, str]],
    sections: Sequence[tuple[str, Sequence[str]]],
    tables: Sequence[tuple[Sequence[str], Sequence[Sequence[Any]]]] = (),
) -> None:
    """Create a deterministic standard-business-brief DOCX."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = f"{title} | Synthetic diligence fixture"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(95, 101, 110)
    footer = section.footer.paragraphs[0]
    footer.text = SYNTHETIC_LABEL
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(7.5)
    footer.runs[0].font.color.rgb = RGBColor(115, 118, 125)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("CONFIDENTIAL - SYNTHETIC FIXTURE")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(122, 90, 0)
    title_p = document.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(23)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle_p = document.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(70, 75, 82)

    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)

    for heading, paragraphs in sections:
        document.add_paragraph(heading, style="Heading 1")
        for text in paragraphs:
            document.add_paragraph(text)

    for headers, rows in tables:
        table = document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table_width = 9360
        first_width = 2700 if len(headers) > 1 else table_width
        remaining = (table_width - first_width) // max(1, len(headers) - 1)
        widths = [first_width] + [remaining] * (len(headers) - 1)
        for index, text in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = str(text)
            _cell_width(cell, widths[index])
            _shade_cell(cell, "E8EEF5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                cells[index].text = str(value)
                _cell_width(cells[index], widths[index])
                cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[index].paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
        document.add_paragraph()

    properties = document.core_properties
    properties.author = "dd-engine fictional fixture generator"
    properties.last_modified_by = "dd-engine fictional fixture generator"
    properties.created = FIXED_DT
    properties.modified = FIXED_DT
    properties.title = title
    properties.subject = SYNTHETIC_LABEL
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    normalize_ooxml(path)


@dataclass(frozen=True, slots=True)
class FormulaCell:
    formula: str
    value: float | int
    style: int = 3


@dataclass(frozen=True, slots=True)
class SheetSpec:
    name: str
    rows: Sequence[Sequence[Any]]
    hidden: bool = False
    hidden_rows: frozenset[int] = field(default_factory=frozenset)
    header_row: int = 3
    freeze_row: int = 3
    currency_columns: frozenset[int] = field(default_factory=frozenset)
    percent_columns: frozenset[int] = field(default_factory=frozenset)
    integer_columns: frozenset[int] = field(default_factory=frozenset)
    warning_cells: frozenset[str] = field(default_factory=frozenset)


def _column_name(number: int) -> str:
    value = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        value = chr(65 + remainder) + value
    return value


def _xlsx_styles() -> bytes:
    return b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
 <numFmts count="3"><numFmt numFmtId="164" formatCode="&#8364;#,##0;[Red](&#8364;#,##0);-"/><numFmt numFmtId="165" formatCode="0.0%"/><numFmt numFmtId="166" formatCode="#,##0;[Red](#,##0);-"/></numFmts>
 <fonts count="5">
  <font><sz val="10"/><name val="Aptos"/><color rgb="FF000000"/></font>
  <font><b/><sz val="15"/><name val="Aptos Display"/><color rgb="FFFFFFFF"/></font>
  <font><b/><sz val="10"/><name val="Aptos"/><color rgb="FFFFFFFF"/></font>
  <font><sz val="10"/><name val="Aptos"/><color rgb="FF0000FF"/></font>
  <font><b/><sz val="10"/><name val="Aptos"/><color rgb="FF000000"/></font>
 </fonts>
 <fills count="5"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill><fill><patternFill patternType="solid"><fgColor rgb="FF0B2545"/><bgColor indexed="64"/></patternFill></fill><fill><patternFill patternType="solid"><fgColor rgb="FF2E74B5"/><bgColor indexed="64"/></patternFill></fill><fill><patternFill patternType="solid"><fgColor rgb="FFFFFF00"/><bgColor indexed="64"/></patternFill></fill></fills>
 <borders count="3"><border><left/><right/><top/><bottom/><diagonal/></border><border><left/><right/><top/><bottom style="thin"><color rgb="FFD9E1EA"/></bottom><diagonal/></border><border><left/><right/><top style="thin"><color rgb="FF000000"/></top><bottom style="double"><color rgb="FF000000"/></bottom><diagonal/></border></borders>
 <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
 <cellXfs count="10">
  <xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>
  <xf numFmtId="0" fontId="1" fillId="2" borderId="0" xfId="0" applyFont="1" applyFill="1"><alignment vertical="center"/></xf>
  <xf numFmtId="0" fontId="2" fillId="3" borderId="1" xfId="0" applyFont="1" applyFill="1" applyBorder="1"><alignment horizontal="center" vertical="center" wrapText="1"/></xf>
  <xf numFmtId="164" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"><alignment horizontal="right"/></xf>
  <xf numFmtId="166" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"><alignment horizontal="right"/></xf>
  <xf numFmtId="165" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"><alignment horizontal="right"/></xf>
  <xf numFmtId="0" fontId="3" fillId="0" borderId="0" xfId="0" applyFont="1"/>
  <xf numFmtId="0" fontId="0" fillId="4" borderId="1" xfId="0" applyFill="1" applyBorder="1"><alignment wrapText="1"/></xf>
  <xf numFmtId="164" fontId="4" fillId="0" borderId="2" xfId="0" applyFont="1" applyBorder="1" applyNumberFormat="1"><alignment horizontal="right"/></xf>
  <xf numFmtId="0" fontId="4" fillId="0" borderId="2" xfId="0" applyFont="1" applyBorder="1"/>
 </cellXfs>
 <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>"""


def _sheet_xml(spec: SheetSpec) -> bytes:
    max_cols = max((len(row) for row in spec.rows), default=1)
    widths: list[float] = []
    for col in range(max_cols):
        lengths = [len(str(row[col])) for row in spec.rows if col < len(row)]
        widths.append(min(42.0, max(11.0, (max(lengths, default=8) + 2) * 1.08)))
    target_total = min(max_cols * 42.0, max(100.0, max_cols * 28.0))
    while sum(widths) < target_total and any(width < 42.0 for width in widths):
        expandable = [index for index, width in enumerate(widths) if width < 42.0]
        increment = (target_total - sum(widths)) / len(expandable)
        for index in expandable:
            widths[index] = min(42.0, widths[index] + increment)
    cols = "".join(
        f'<col min="{index}" max="{index}" width="{width:.1f}" customWidth="1"/>'
        for index, width in enumerate(widths, start=1)
    )
    row_xml: list[str] = []
    for row_number, row in enumerate(spec.rows, start=1):
        hidden = ' hidden="1"' if row_number in spec.hidden_rows else ""
        height = ' ht="27" customHeight="1"' if row_number in {1, spec.header_row} else ""
        cells: list[str] = []
        for col_number, raw in enumerate(row, start=1):
            if raw is None:
                continue
            ref = f"{_column_name(col_number)}{row_number}"
            style = 0
            if row_number == 1:
                style = 1
            elif row_number == spec.header_row:
                style = 2
            elif ref in spec.warning_cells:
                style = 7
            elif isinstance(raw, FormulaCell):
                style = raw.style
            elif col_number in spec.currency_columns and isinstance(raw, int | float):
                style = 3
            elif col_number in spec.percent_columns and isinstance(raw, int | float):
                style = 5
            elif col_number in spec.integer_columns and isinstance(raw, int | float):
                style = 4
            if isinstance(raw, FormulaCell):
                cells.append(
                    f'<c r="{ref}" s="{style}"><f>{escape(raw.formula)}</f><v>{raw.value}</v></c>'
                )
            elif isinstance(raw, bool):
                cells.append(f'<c r="{ref}" s="{style}" t="b"><v>{int(raw)}</v></c>')
            elif isinstance(raw, int | float):
                cells.append(f'<c r="{ref}" s="{style}"><v>{raw}</v></c>')
            else:
                cells.append(
                    f'<c r="{ref}" s="{style}" t="inlineStr"><is><t>{escape(str(raw))}</t></is></c>'
                )
        row_xml.append(f'<row r="{row_number}"{hidden}{height}>{"".join(cells)}</row>')
    merges = (
        f'<mergeCells count="1"><mergeCell ref="A1:{_column_name(max_cols)}1"/></mergeCells>'
        if max_cols > 1
        else ""
    )
    filter_xml = ""
    if len(spec.rows) >= spec.header_row:
        filter_xml = (
            f'<autoFilter ref="A{spec.header_row}:{_column_name(max_cols)}{len(spec.rows)}"/>'
        )
    pane = ""
    if spec.freeze_row:
        pane = f'<pane ySplit="{spec.freeze_row}" topLeftCell="A{spec.freeze_row + 1}" activePane="bottomLeft" state="frozen"/>'
    xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
 <sheetViews><sheetView showGridLines="0" workbookViewId="0">{pane}</sheetView></sheetViews>
 <sheetFormatPr defaultRowHeight="15"/><cols>{cols}</cols><sheetData>{"".join(row_xml)}</sheetData>
 {filter_xml}{merges}<pageMargins left="0.35" right="0.35" top="0.5" bottom="0.5" header="0.2" footer="0.2"/>
</worksheet>"""
    return xml.encode("utf-8")


def write_xlsx(path: Path, sheets: Sequence[SheetSpec]) -> None:
    """Write a deterministic, styled, valid XLSX using direct OOXML."""

    if not sheets:
        raise ValueError("an XLSX requires at least one sheet")
    path.parent.mkdir(parents=True, exist_ok=True)
    sheet_nodes = []
    rel_nodes = []
    content_nodes = []
    members: dict[str, bytes] = {}
    for index, spec in enumerate(sheets, start=1):
        state = ' state="hidden"' if spec.hidden else ""
        sheet_nodes.append(
            f'<sheet name={quoteattr(spec.name)} sheetId="{index}"{state} r:id="rId{index}"/>'
        )
        rel_nodes.append(
            f'<Relationship Id="rId{index}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{index}.xml"/>'
        )
        content_nodes.append(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        )
        members[f"xl/worksheets/sheet{index}.xml"] = _sheet_xml(spec)
    rel_nodes.append(
        '<Relationship Id="rIdStyles" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
    )
    members["[Content_Types].xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        + "".join(content_nodes)
        + "</Types>"
    ).encode()
    members["_rels/.rels"] = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
 <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
 <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
 <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>"""
    members["xl/workbook.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<bookViews><workbookView activeTab="0"/></bookViews><sheets>{"".join(sheet_nodes)}</sheets>'
        '<calcPr calcId="191029" fullCalcOnLoad="1" forceFullCalc="1"/></workbook>'
    ).encode()
    members["xl/_rels/workbook.xml.rels"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(rel_nodes)
        + "</Relationships>"
    ).encode()
    members["xl/styles.xml"] = _xlsx_styles()
    members["docProps/core.xml"] = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>Synthetic diligence workbook</dc:title><dc:creator>dd-engine fictional fixture generator</dc:creator><dc:subject>FICTIONAL SYNTHETIC DATA</dc:subject><dcterms:created xsi:type="dcterms:W3CDTF">2026-08-31T09:00:00Z</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">2026-08-31T09:00:00Z</dcterms:modified></cp:coreProperties>"""
    members["docProps/app.xml"] = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>dd-engine synthetic generator</Application><DocSecurity>0</DocSecurity><ScaleCrop>false</ScaleCrop></Properties>"""
    write_stable_zip(path, members)


def write_csv_bytes_with_xlsx_extension(path: Path, rows: Sequence[Sequence[Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    encoded_rows = []
    for row in rows:
        fields = []
        for value in row:
            text = str(value).replace('"', '""')
            fields.append(f'"{text}"')
        encoded_rows.append(",".join(fields))
    path.write_text("\n".join(encoded_rows) + "\n", encoding="utf-8", newline="")


def write_phone_photo(
    path: Path,
    *,
    bank_name: str,
    reference: str,
    amount: int,
    text: Sequence[str],
    seed: int,
    angle: float,
) -> None:
    """Create a genuine JPEG with desk, rotated letter, skew and sensor noise."""

    rng = random.Random(seed)
    desk = Image.new("RGB", (1600, 1200), (118, 82, 54))
    desk_draw = ImageDraw.Draw(desk)
    for y in range(0, 1200, 9):
        tone = 105 + int(10 * math.sin(y / 41)) + rng.randint(-4, 4)
        desk_draw.line((0, y, 1600, y), fill=(tone + 18, tone, max(40, tone - 24)), width=8)
    letter = Image.new("RGB", (920, 1180), (248, 246, 238))
    draw = ImageDraw.Draw(letter)
    draw.rectangle((55, 46, 865, 1130), outline=(176, 174, 167), width=2)
    draw.text((85, 85), bank_name, font=_font(38, bold=True), fill=(22, 48, 78))
    draw.text((85, 145), "CREDIT FACILITIES", font=_font(18, bold=True), fill=(130, 92, 20))
    draw.text((85, 195), SYNTHETIC_LABEL, font=_font(15, bold=True), fill=(145, 35, 35))
    draw.text((85, 255), f"Reference: {reference}", font=_font(22), fill=(38, 38, 36))
    draw.text((85, 300), "31 August 2026", font=_font(22), fill=(38, 38, 36))
    y = 375
    for paragraph in text:
        for line in _wrapped_lines(paragraph, 58):
            draw.text((90, y), line, font=_font(21), fill=(50, 49, 45))
            y += 31
        y += 20
    draw.text(
        (90, 880), f"Facility amount: EUR {amount:,}", font=_font(25, bold=True), fill=(22, 48, 78)
    )
    draw.text((90, 980), "Yours faithfully,", font=_font(21), fill=(50, 49, 45))
    draw.line((90, 1060, 390, 1020), fill=(30, 58, 90), width=4)
    draw.text((90, 1070), "Synthetic Credit Manager", font=_font(18), fill=(50, 49, 45))

    # A non-uniform resize creates a mild perspective skew before rotation.
    warped = letter.transform(
        (980, 1120),
        Image.Transform.QUAD,
        (35, 30, 5, 1115, 915, 1160, 885, 0),
        resample=Image.Resampling.BICUBIC,
    )
    warped_rgba = warped.convert("RGBA").rotate(
        angle,
        resample=Image.Resampling.BICUBIC,
        expand=True,
        fillcolor=(0, 0, 0, 0),
    )
    mask = warped_rgba.getchannel("A")
    shadow_mask = mask.filter(ImageFilter.GaussianBlur(18))
    x = 270 + rng.randint(-30, 30)
    y_pos = 35 + rng.randint(-15, 15)
    desk.paste((28, 22, 18), (x + 20, y_pos + 24), shadow_mask)
    desk.paste(warped_rgba.convert("RGB"), (x, y_pos), mask)
    pixels = desk.load()
    for _ in range(42000):
        px = rng.randrange(desk.width)
        py = rng.randrange(desk.height)
        r, g, b = pixels[px, py]
        delta = rng.choice((-3, -2, -1, 1, 2, 3))
        pixels[px, py] = tuple(max(0, min(255, channel + delta)) for channel in (r, g, b))
    path.parent.mkdir(parents=True, exist_ok=True)
    desk.save(path, format="JPEG", quality=83, optimize=False, progressive=False, dpi=(96, 96))


def write_cro_screenshot(path: Path, *, company: str, registration_id: str) -> None:
    image = Image.new("RGB", (1440, 900), (246, 248, 251))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1440, 74), fill=(39, 62, 92))
    draw.circle((35, 37), 11, fill=(233, 90, 80))
    draw.circle((70, 37), 11, fill=(236, 184, 72))
    draw.circle((105, 37), 11, fill=(91, 190, 105))
    draw.rounded_rectangle((150, 18, 1110, 57), radius=14, fill=(255, 255, 255))
    draw.text((180, 27), "cro-fixture.invalid/search/company", font=_font(18), fill=(74, 79, 88))
    draw.text(
        (90, 130),
        "Companies Registration Office - search preview",
        font=_font(31, bold=True),
        fill=(29, 49, 77),
    )
    draw.text((90, 185), SYNTHETIC_LABEL, font=_font(19, bold=True), fill=(150, 38, 38))
    draw.rounded_rectangle(
        (90, 245, 1345, 760), radius=12, fill=(255, 255, 255), outline=(205, 212, 222), width=2
    )
    rows = [
        ("Company name", company),
        ("Synthetic registration", registration_id),
        ("Status", "Normal - fictional fixture"),
        ("Next annual return", "2027-03-14"),
        ("Registered office", "14 Alder Quay, Kilnmore Business Park, Dublin D99 SYN4"),
        ("Important", "Screenshot only - not a certified extract"),
    ]
    y = 290
    for label, value in rows:
        draw.text((135, y), label, font=_font(20, bold=True), fill=(55, 72, 95))
        draw.text((420, y), value, font=_font(20), fill=(45, 48, 54))
        draw.line((125, y + 43, 1310, y + 43), fill=(226, 230, 236), width=2)
        y += 73
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=False)


def bytes_from_writer(writer: Any, suffix: str, *args: Any, **kwargs: Any) -> bytes:
    """Run a path-oriented writer in a temporary directory and return its bytes."""

    import tempfile

    with tempfile.TemporaryDirectory(prefix="dd-synthetic-member-") as temporary:
        path = Path(temporary) / f"member{suffix}"
        writer(path, *args, **kwargs)
        return path.read_bytes()


def rows_total(rows: Iterable[Sequence[Any]], index: int) -> float:
    return float(sum(float(row[index]) for row in rows))
