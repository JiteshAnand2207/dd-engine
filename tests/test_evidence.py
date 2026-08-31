from __future__ import annotations

import hashlib
import json
import shutil
from dataclasses import replace
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image
from reportlab.pdfgen import canvas

from dd_engine.config import EngineConfig, load_config
from dd_engine.evidence import (
    Calculation,
    CitationValidator,
    Claim,
    Evidence,
    build_evidence_foundation,
)
from dd_engine.evidence.calculations import recompute_calculation
from dd_engine.evidence.models import JsonObject
from dd_engine.extraction import extract_run
from dd_engine.inventory import RegisterLimits, register_room
from dd_engine.runs import create_run, load_manifest
from dd_engine.state import mark_stage_awaiting_input, start_stage

LIMITS = RegisterLimits(
    max_archive_members=100,
    max_archive_total_uncompressed_bytes=16 * 1024 * 1024,
    max_archive_member_uncompressed_bytes=4 * 1024 * 1024,
)


def _config_without_ocr(root: Path) -> EngineConfig:
    config = load_config(cwd=root)
    return replace(config, extraction=replace(config.extraction, optional_ocr=False))


def _write_pdf(path: Path) -> None:
    document = canvas.Canvas(str(path))
    document.drawString(72, 720, "Page one source fact: reported amount is EUR 15.")
    document.showPage()
    document.drawString(72, 720, "Page two source fact.")
    document.save()


def _write_workbook(path: Path, values: tuple[int, int]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet["A1"] = "Metric"
    sheet["B1"] = "Value"
    sheet["A2"] = "Input A"
    sheet["B2"] = values[0]
    sheet["A3"] = "Input B"
    sheet["B3"] = values[1]
    workbook.save(path)


def _build_room(room: Path) -> None:
    room.mkdir()
    _write_pdf(room / "Citation.pdf")
    shutil.copyfile(room / "Citation.pdf", room / "Duplicate_One.pdf")
    shutil.copyfile(room / "Citation.pdf", room / "Duplicate_Two.pdf")
    _write_workbook(room / "Metrics_Original.xlsx", (8, 4))
    _write_workbook(room / "Metrics_Rev2.xlsx", (10, 5))
    document = Document()
    document.add_paragraph("Paragraph one source fact.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Table fact"
    table.cell(0, 1).text = "Supported"
    document.save(room / "Memo.docx")
    (room / "Rows.csv").write_text("Name,Value\nAlpha,15\n", encoding="utf-8")
    Image.new("RGB", (40, 20), "white").save(room / "Photo.png", format="PNG")


@pytest.fixture(scope="module")
def evidence_run(tmp_path_factory: pytest.TempPathFactory) -> Path:
    root = tmp_path_factory.mktemp("evidence-ready")
    room = root / "room"
    _build_room(room)
    config = _config_without_ocr(root)
    run_path = create_run(config, runs_root=root / "runs")
    register_room(run_path, room, LIMITS)
    extract_run(run_path, room, config)
    return run_path


def _register(run_path: Path) -> list[JsonObject]:
    payload = json.loads(
        (run_path / "source_register" / "source_register.json").read_text(encoding="utf-8")
    )
    return payload["sources"]


def _source(run_path: Path, suffix: str) -> JsonObject:
    return next(item for item in _register(run_path) if item["relative_path"].endswith(suffix))


def _units(run_path: Path, source_id: str) -> list[JsonObject]:
    return [
        json.loads(line)
        for line in (run_path / "extracts" / "extracted_units.jsonl")
        .read_text(encoding="utf-8")
        .splitlines()
        if line and json.loads(line)["source_id"] == source_id
    ]


def _unit(
    run_path: Path,
    source: JsonObject,
    unit_type: str,
    **locator_values: object,
) -> JsonObject:
    for item in _units(run_path, str(source["source_id"])):
        if item["unit_type"] != unit_type:
            continue
        if all(item["locator"].get(key) == value for key, value in locator_values.items()):
            return item
    raise AssertionError(f"unit not found: {source['source_id']} {unit_type} {locator_values}")


def _claim(run_path: Path, claim_id: str, *, required_sources: int = 1) -> JsonObject:
    return Claim(
        run_id=run_path.name,
        claim_id=claim_id,
        statement=f"Material source-backed statement {claim_id}",
        claim_type="fact",
        workstream="financial",
        materiality="high",
        confidence=0.9,
        status="supported",
        required_independent_sources=required_sources,
    ).as_record()


def _evidence(
    run_path: Path,
    claim_id: str,
    evidence_id: str,
    source: JsonObject,
    unit: JsonObject,
    *,
    acknowledge_supersession: bool = False,
) -> JsonObject:
    content = unit["content"]
    text = content.get("text") if isinstance(content.get("text"), str) else None
    value: Any = None
    if text is None:
        for key in ("value", "source_value", "cached_value"):
            if content.get(key) is not None:
                value = content[key]
                break
        if value is None:
            value = content
    return Evidence(
        run_id=run_path.name,
        evidence_id=evidence_id,
        claim_id=claim_id,
        source_id=str(source["source_id"]),
        source_checksum=str(source["sha256"]),
        exact_locator=dict(unit["locator"]),
        relationship="supporting",
        extraction_confidence=float(unit["confidence"]),
        source_version_status=str(source["probable_version_status"]),
        extracted_value=value,
        extracted_text=text,
        extracted_unit_ids=(str(unit["unit_id"]),),
        supersession_acknowledged=acknowledge_supersession,
    ).as_record()


def _record_sets(
    *,
    claims: list[JsonObject] | None = None,
    evidence: list[JsonObject] | None = None,
    calculations: list[JsonObject] | None = None,
) -> dict[str, list[JsonObject]]:
    return {
        "calculations": calculations or [],
        "claims": claims or [],
        "contradictions": [],
        "evidence": evidence or [],
        "gaps": [],
        "issues": [],
    }


def test_foundation_writes_all_outputs_without_starting_analysis(evidence_run: Path) -> None:
    first = build_evidence_foundation(evidence_run)
    second = build_evidence_foundation(evidence_run)
    _, manifest = load_manifest(evidence_run)

    assert first.validation_passed is True
    assert second.reused is True
    assert manifest["stages"]["analyse"]["state"] == "not_started"
    assert first.summary["material_claim_coverage"] is None
    for filename in (
        "claims.jsonl",
        "evidence.jsonl",
        "calculations.jsonl",
        "contradictions.jsonl",
        "gaps.jsonl",
        "issues.jsonl",
        "citation_validation.json",
        "evidence_coverage.md",
    ):
        assert (evidence_run / "evidence" / filename).is_file()
    gaps = [
        json.loads(line)
        for line in (evidence_run / "evidence" / "gaps.jsonl")
        .read_text(encoding="utf-8")
        .splitlines()
    ]
    assert any(item["origin"] == "foundation_vision" for item in gaps)


def test_available_answers_are_preserved_and_unanswered_intake_remains_a_gap(
    evidence_run: Path, tmp_path: Path
) -> None:
    run_path = tmp_path / evidence_run.name
    shutil.copytree(evidence_run, run_path)
    start_stage(
        run_path,
        "intake",
        input_checksum=hashlib.sha256(b"phase7-intake-fixture").hexdigest(),
    )
    mark_stage_awaiting_input(run_path, "intake")
    questions = {
        "questions": [
            {
                "decision_potentially_affected": ["price"],
                "exact_question": "Provide the supported amount.",
                "priority": "critical",
                "question_id": "INT-R1-001",
                "structured_gap": {
                    "description": "The amount is not source-supported.",
                    "gap_id": "GAP-AMOUNT",
                },
                "supporting_source_ids": [],
            },
            {
                "decision_potentially_affected": ["go_no_go"],
                "exact_question": "Provide the missing contract.",
                "priority": "high",
                "question_id": "INT-R1-002",
                "structured_gap": {
                    "description": "The contract is absent.",
                    "gap_id": "GAP-CONTRACT",
                },
                "supporting_source_ids": [],
            },
        ],
        "round_number": 1,
        "run_id": run_path.name,
        "status": "awaiting_input",
    }
    answers = {
        "answers": [
            {
                "ambiguity": ["Reply points elsewhere and remains unresolved."],
                "normalised_interpretation": {
                    "kind": "cross_reference",
                    "value": "see finance 2.1",
                },
                "provenance": {"source_kind": "deal_lead_answer_file"},
                "question_id": "INT-R1-001",
                "resolution_status": "narrowed",
                "verbatim_answer": "see finance 2.1",
                "verbatim_answer_sha256": hashlib.sha256(b"see finance 2.1").hexdigest(),
            }
        ],
        "round_number": 1,
        "run_id": run_path.name,
        "status": "ingested",
    }
    (run_path / "intake" / "round_1_questions.json").write_text(
        json.dumps(questions), encoding="utf-8"
    )
    (run_path / "intake" / "round_1_answers.json").write_text(json.dumps(answers), encoding="utf-8")

    build_evidence_foundation(run_path)
    gaps = {
        item["gap_id"]: item
        for item in (
            json.loads(line)
            for line in (run_path / "evidence" / "gaps.jsonl")
            .read_text(encoding="utf-8")
            .splitlines()
        )
    }
    _, manifest = load_manifest(run_path)

    assert gaps["GAP-AMOUNT"]["status"] == "narrowed"
    assert gaps["GAP-AMOUNT"]["answer_provenance"]["verbatim_answer"] == ("see finance 2.1")
    assert gaps["GAP-CONTRACT"]["status"] == "open"
    assert gaps["GAP-CONTRACT"]["answer_provenance"] is None
    assert manifest["stages"]["intake"]["state"] == "awaiting_input"
    assert manifest["stages"]["analyse"]["state"] == "not_started"


def test_valid_pdf_spreadsheet_docx_csv_and_image_citations(evidence_run: Path) -> None:
    fixtures = [
        ("Citation.pdf", "pdf_page", {"page_number": 1}),
        ("Metrics_Rev2.xlsx", "spreadsheet_cell", {"cell": "B2"}),
        ("Memo.docx", "docx_paragraph", {"paragraph_index": 1}),
        ("Memo.docx", "docx_table_cell", {"table_index": 1, "row_index": 1, "cell_index": 1}),
        ("Rows.csv", "csv_cell", {"row_index": 2, "column_index": 2}),
        ("Photo.png", "image_metadata", {"image_number": 1}),
    ]
    claims: list[JsonObject] = []
    citations: list[JsonObject] = []
    for index, (suffix, unit_type, locator) in enumerate(fixtures, start=1):
        source = _source(evidence_run, suffix)
        unit = _unit(evidence_run, source, unit_type, **locator)
        claim_id = f"CLM-VALID-{index}"
        claims.append(_claim(evidence_run, claim_id))
        citations.append(_evidence(evidence_run, claim_id, f"EVD-VALID-{index}", source, unit))

    report = CitationValidator(evidence_run).validate(
        _record_sets(claims=claims, evidence=citations)
    )

    assert report["status"] == "passed"
    assert report["summary"]["failed_citation_count"] == 0
    assert report["summary"]["material_claim_coverage"] == 1.0


@pytest.mark.parametrize(
    ("suffix", "unit_type", "unit_locator", "invalid_locator", "expected_code"),
    [
        (
            "Citation.pdf",
            "pdf_page",
            {"page_number": 1},
            {"type": "pdf_page", "page_number": 99},
            "pdf_page_not_found",
        ),
        (
            "Metrics_Rev2.xlsx",
            "spreadsheet_cell",
            {"cell": "B2"},
            {"type": "spreadsheet_cell", "sheet": "Missing", "cell": "B2", "range": "B2"},
            "spreadsheet_sheet_not_found",
        ),
        (
            "Memo.docx",
            "docx_paragraph",
            {"paragraph_index": 1},
            {"type": "docx_paragraph", "paragraph_index": 999},
            "docx_paragraph_not_found",
        ),
    ],
)
def test_invalid_format_native_citations_fail(
    evidence_run: Path,
    suffix: str,
    unit_type: str,
    unit_locator: dict[str, object],
    invalid_locator: JsonObject,
    expected_code: str,
) -> None:
    source = _source(evidence_run, suffix)
    unit = _unit(evidence_run, source, unit_type, **unit_locator)
    claim = _claim(evidence_run, "CLM-INVALID")
    citation = _evidence(evidence_run, claim["claim_id"], "EVD-INVALID", source, unit)
    citation["exact_locator"] = invalid_locator

    report = CitationValidator(evidence_run).validate(
        _record_sets(claims=[claim], evidence=[citation])
    )

    assert report["status"] == "failed"
    assert expected_code in {error["code"] for error in report["failed_citations"][0]["errors"]}


def test_exact_duplicates_do_not_count_as_independent_corroboration(
    evidence_run: Path,
) -> None:
    sources = [
        item
        for item in _register(evidence_run)
        if item["relative_path"].endswith(("Duplicate_One.pdf", "Duplicate_Two.pdf"))
    ]
    assert len(sources) == 2
    assert sources[0]["duplicate_group"] == sources[1]["duplicate_group"]
    claim = _claim(evidence_run, "CLM-DUPLICATE", required_sources=2)
    citations = [
        _evidence(
            evidence_run,
            claim["claim_id"],
            f"EVD-DUPLICATE-{index}",
            source,
            _unit(evidence_run, source, "pdf_page", page_number=1),
        )
        for index, source in enumerate(sources, start=1)
    ]

    report = CitationValidator(evidence_run).validate(
        _record_sets(claims=[claim], evidence=citations)
    )

    assert report["status"] == "failed"
    assert report["claim_results"][0]["supporting_citation_count"] == 2
    assert report["claim_results"][0]["independent_supporting_source_count"] == 1
    assert report["summary"]["duplicate_corroboration_exclusion_count"] == 1


def test_superseded_source_requires_explicit_acknowledgement(evidence_run: Path) -> None:
    source = next(
        item
        for item in _register(evidence_run)
        if item["probable_version_status"] == "potentially_superseded"
    )
    unit = _units(evidence_run, str(source["source_id"]))[0]
    claim = _claim(evidence_run, "CLM-SUPERSEDED")
    citation = _evidence(evidence_run, claim["claim_id"], "EVD-SUPERSEDED", source, unit)

    silent = CitationValidator(evidence_run).validate(
        _record_sets(claims=[claim], evidence=[citation])
    )
    citation["supersession_acknowledged"] = True
    acknowledged = CitationValidator(evidence_run).validate(
        _record_sets(claims=[claim], evidence=[citation])
    )

    assert "silently_superseded_source" in {
        error["code"] for error in silent["failed_citations"][0]["errors"]
    }
    assert acknowledged["status"] == "passed"
    assert acknowledged["citation_results"][0]["warnings"][0]["code"] == (
        "superseded_source_acknowledged"
    )


def _source_input(source: JsonObject, unit: JsonObject, input_id: str) -> JsonObject:
    value = unit["content"]["source_value"]
    return {
        "currency": "EUR",
        "extracted_unit_ids": [unit["unit_id"]],
        "input_id": input_id,
        "locator": unit["locator"],
        "missing": False,
        "normalized_value": value,
        "period": "not_applicable",
        "reported_value": value,
        "sign_convention": "positive_as_reported",
        "source_checksum": source["sha256"],
        "source_id": source["source_id"],
        "source_version_status": source["probable_version_status"],
        "supersession_acknowledged": False,
        "unit": "EUR",
    }


def _calculation(evidence_run: Path) -> JsonObject:
    source = _source(evidence_run, "Metrics_Rev2.xlsx")
    input_a = _unit(evidence_run, source, "spreadsheet_cell", cell="B2")
    input_b = _unit(evidence_run, source, "spreadsheet_cell", cell="B3")
    return Calculation(
        run_id=evidence_run.name,
        calculation_id="CALC-VALID",
        description="Recompute the two reported workbook inputs.",
        source_inputs=(
            _source_input(source, input_a, "input_a"),
            _source_input(source, input_b, "input_b"),
        ),
        units="EUR",
        currency="EUR",
        normalisation={
            "currency": "EUR as reported; no conversion",
            "period": "not applicable; point values",
            "sign": "positive values retained",
            "units": "EUR, no scaling",
        },
        formula={"expression": "input_a + input_b", "version": "sum-v1"},
        result={"reported_value": 14, "recomputed_value": 15, "variance": 1},
        rounding={"decimal_places": 2, "mode": "none"},
        independent_recomputation_status="verified",
        calculation_method="deterministic",
    ).as_record()


def test_calculation_keeps_reported_and_recomputed_values_separate(
    evidence_run: Path,
) -> None:
    calculation = _calculation(evidence_run)
    report = CitationValidator(evidence_run).validate(_record_sets(calculations=[calculation]))

    assert calculation["result"] == {
        "reported_value": 14,
        "recomputed_value": 15,
        "variance": 1,
    }
    assert report["status"] == "passed"
    assert report["calculation_results"][0]["recomputation"]["status"] == "verified"


def test_invalid_calculation_citation_fails(evidence_run: Path) -> None:
    calculation = _calculation(evidence_run)
    calculation["source_inputs"][0]["source_checksum"] = "0" * 64

    report = CitationValidator(evidence_run).validate(_record_sets(calculations=[calculation]))

    assert report["status"] == "failed"
    assert report["summary"]["failed_citation_count"] == 1
    assert report["failed_citations"][0]["errors"][0]["code"] == ("source_checksum_mismatch")


def test_missing_calculation_input_is_not_assumed_zero(evidence_run: Path) -> None:
    calculation = _calculation(evidence_run)
    calculation["source_inputs"] = list(calculation["source_inputs"])
    calculation["source_inputs"][1] = {
        "input_id": "input_b",
        "missing": True,
        "missing_reason": "Requested schedule was not supplied.",
        "normalized_value": None,
        "reported_value": None,
    }
    calculation["result"]["recomputed_value"] = None
    calculation["result"]["variance"] = None
    calculation["independent_recomputation_status"] = "blocked_missing_inputs"

    recomputation = recompute_calculation(calculation)
    report = CitationValidator(evidence_run).validate(_record_sets(calculations=[calculation]))

    assert recomputation["status"] == "blocked_missing_inputs"
    assert recomputation["errors"] == []
    assert report["status"] == "passed"
